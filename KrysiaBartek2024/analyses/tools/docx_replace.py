"""
This script reads a DOCX file, replaces specific strings with symbols, and colors them.
The output is saved as a new DOCX file.
"""

import os
# pylint: disable=E0401
from docx import Document
from docx.shared import RGBColor

OUTPUT_DIR = './tmp'

INPUT_FILE = os.path.join(OUTPUT_DIR, 'to_docx.docx')
OUTPUT_FILE = os.path.join(OUTPUT_DIR, 'ZACZEK.docx')

replacement_map = {
    '!c!': ('♣', RGBColor(0, 128, 0)),     # Clubs (green)
    '!s!': ('♠', RGBColor(0, 0, 255)),     # Spades (blue)
    '!h!': ('♥', RGBColor(255, 0, 0)),     # Hearts (red)
    '!d!': ('♦', RGBColor(255, 165, 0))    # Diamonds (orange)
}

doc = Document(INPUT_FILE)

def replace_and_color_run(paragraph, replacements):
    """
    Replace and color the symbols in the given paragraph.
    """
    # Get the entire text of the paragraph
    original_text = paragraph.text

    # If any replacement string exists in the paragraph text, we process it
    if any(old_str in original_text for old_str in replacements):
        paragraph.clear()

        while original_text:
            # Find the first occurrence of any replacement string
            next_replacement = min(
                ((old_str, original_text.find(old_str)) for old_str \
                    in replacements if old_str in original_text),
                key=lambda x: x[1],
                default=(None, -1)
            )

            old_str, index = next_replacement

            # If no more replacements found, add the rest of the text and break
            if index == -1:
                paragraph.add_run(original_text)  # Add remaining text
                break

            # Add the text before the replacement symbol
            if index > 0:
                paragraph.add_run(original_text[:index])

            # Add the replacement symbol with color
            symbol, color = replacements[old_str]
            colored_run = paragraph.add_run(symbol)
            colored_run.font.color.rgb = color

            # Update the text to be the part after the current replacement
            original_text = original_text[index + len(old_str):]

# Iterate over paragraphs and process their text
for para in doc.paragraphs:
    replace_and_color_run(para, replacement_map)

doc.save(OUTPUT_FILE)

print(f"Replacement and coloring complete. Saved as {OUTPUT_FILE}.")
